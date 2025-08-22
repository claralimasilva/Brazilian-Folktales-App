#!/usr/bin/env python3
# Script to debug DOCX content and understand quiz format

from docx import Document
import re

def debug_docx_content(docx_file):
    """Debug DOCX content to understand structure"""
    try:
        doc = Document(docx_file)
        print(f"=== DEBUGGING {docx_file} ===")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print("="*50)
        
        quiz_patterns = []
        current_chapter = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Look for chapter markers
            if 'chapter' in text.lower() or 'capítulo' in text.lower():
                current_chapter = text
                print(f"\n📖 CHAPTER FOUND (line {i}): {text}")
                continue
            
            # Look for quiz patterns
            quiz_indicators = ['quiz', 'question', 'pergunta', 'resposta', 'answer', 'a)', 'b)', 'c)', 'd)', '1.', '2.', '3.']
            
            if any(indicator in text.lower() for indicator in quiz_indicators):
                quiz_patterns.append({
                    'line': i,
                    'chapter': current_chapter,
                    'text': text,
                    'style': paragraph.style.name if paragraph.style else 'No style'
                })
                print(f"🎯 POTENTIAL QUIZ (line {i}): {text[:100]}...")
                print(f"   Style: {paragraph.style.name if paragraph.style else 'No style'}")
        
        print(f"\n=== QUIZ ANALYSIS ===")
        print(f"Found {len(quiz_patterns)} potential quiz elements")
        
        if quiz_patterns:
            print("\nAll quiz patterns found:")
            for pattern in quiz_patterns:
                print(f"Line {pattern['line']}: {pattern['text']}")
        
        # Look for specific patterns
        print("\n=== PATTERN ANALYSIS ===")
        question_patterns = []
        answer_patterns = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Question patterns
            if re.search(r'\?', text) or 'question' in text.lower() or 'pergunta' in text.lower():
                question_patterns.append(text)
            
            # Answer patterns  
            if re.search(r'^[a-d]\)', text.lower()) or re.search(r'^\d+\.', text):
                answer_patterns.append(text)
        
        print(f"Question patterns found: {len(question_patterns)}")
        for q in question_patterns[:5]:  # Show first 5
            print(f"  Q: {q}")
            
        print(f"\nAnswer patterns found: {len(answer_patterns)}")
        for a in answer_patterns[:10]:  # Show first 10
            print(f"  A: {a}")
            
    except Exception as e:
        print(f"Error reading DOCX: {e}")

if __name__ == "__main__":
    debug_docx_content('BrazilianFolktales.docx')
    print("\n" + "="*50)
    debug_docx_content('Brazilian Folktales.docx')
